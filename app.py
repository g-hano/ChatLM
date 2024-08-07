from flask import Flask, request, render_template, jsonify
from flask import stream_with_context, Response, send_from_directory
from werkzeug.utils import secure_filename
import os
import requests
from llama_index.core import Settings
from HybridRetriever import HybridRetriever
from ChatEngine import ChatEngine
from configs import *
import csv
import warnings
warnings.simplefilter("ignore", UserWarning)
warnings.simplefilter("ignore", FutureWarning)
import logging
logging.basicConfig(level=logging.INFO)
from llama_index.retrievers.bm25 import BM25Retriever 
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core import VectorStoreIndex, Document
from llama_index.core import Settings
from llama_index.core.node_parser import SentenceSplitter
import fitz
from docx import Document as DocxDocument
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
import json

def process_file(file):
    file_extension = file.split(".")[-1].lower()

    if file_extension == 'txt':
        with open(file, 'r', encoding='utf-8') as f:
            text = f.read()

    elif file_extension == 'csv':
        with open(file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            text = '\n'.join(','.join(row) for row in reader)

    elif file_extension == 'pdf':
        pdf_document = fitz.open(file, filetype=file_extension)
        text = ""
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
        pdf_document.close()
        
    elif file_extension == 'docx':
        docx_document = DocxDocument(file)
        text = ""
        for paragraph in docx_document.paragraphs:
            text += paragraph.text + "\n"

    return [Document(text=text)]

def send_request_2_llm(prompt: str):
    url = "http://localhost:8000/generate"
    if len(prompt) > 27_000: # model-len is set to 27k on vllm.entrypoints.api_server
        prompt = prompt[:27_000]
    payload = {
        "prompt": prompt,
        "stream": True,
        "min_tokens": 256,
        "max_tokens": 1024
    }
    last_response_len = len(prompt)
    with requests.post(url, json=payload, stream=True) as response:
        if response.status_code == 200:
            for line in response.iter_lines():
                if line:
                    try:
                        chunk = json.loads(line.decode('utf-8'))
                        yield chunk.text[last_response_len:]
                        last_response_len = len(chunk.text[last_response:])#json.dumps(chunk) 
                    except json.JSONDecodeError:
                        print(f"Failed to decode JSON: {line.decode('utf-8')}")
        else:
            yield json.dumps({"text":f"Error: Received status code {response.status_code}"})


def process_and_respond(file, question):
    documents = process_file(file)
    
    text_splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    vector_index = VectorStoreIndex.from_documents(
        documents, transformations=[text_splitter], embed_model=Settings.embed_model, show_progress=True
    )
    bm25_retriever = BM25Retriever(nodes=documents, similarity_top_k=TOP_K, tokenizer=text_splitter.split_text)
    vector_retriever = VectorIndexRetriever(index=vector_index, similarity_top_k=TOP_K)
    hybrid_retriever = HybridRetriever(bm25_retriever=bm25_retriever, vector_retriever=vector_retriever)
    chat_engine = ChatEngine(hybrid_retriever)
    chat_history = chat_engine.ask_question(question)  # returns modified text with relevant documents + question
    return len(chat_history), send_request_2_llm(chat_history)



if __name__ == '__main__':
    app = Flask(__name__, static_folder="static", static_url_path="/static")
    app.config['UPLOAD_FOLDER'] = 'uploads'

    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

    embedding = HuggingFaceEmbedding(
        model_name=EMBEDDING_NAME,
        device="cuda:2",
        trust_remote_code=True,
        )
    Settings.embed_model = embedding
    
    @app.route('/', methods=['GET', 'POST'])
    def home():
        if request.method == 'POST':
           return upload_file()
        return render_template('index.html')
    
    @app.route('/upload', methods=['POST'])
    def upload_file():
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if file:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            question = request.form.get('question', '')
            
            def generate():
                system_prompt_length, response_generator = process_and_respond(file_path, question)
                logging.debug(f"System prompt length: {system_prompt_length}")
                last_response = ""
                try:
                   for chunk in response_generator:
                       try:
                          # Attempt to parse the chunk as JSON
                          json_chunk = json.loads(chunk)
                          if "text" in json_chunk:
                             full_text = json_chunk["text"]
                             if len(full_text) > system_prompt_length:
                                new_content = full_text[system_prompt_length:]
                                if len(new_content) > len(last_response):
                                   yield new_content[len(last_response):] + '\n'
                                   last_response = new_content
                          else:
                              logging.debug("Received JSON chunk without 'text' key")
                       except json.JSONDecodeError:
                           # If it's not JSON, process it as before
                           if len(chunk) > system_prompt_length:
                              new_content = chunk[system_prompt_length:]
                              if len(new_content) > len(last_response):
                                 yield new_content[len(last_response):] + '\n'
                                 last_response = new_content
                           else:
                               logging.debug("Chunk not longer than system prompt")
                finally:
                    os.remove(file_path)

            return Response(stream_with_context(generate()), content_type='text/plain')
    app.run(host='0.0.0.0', port=5000)

