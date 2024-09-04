import csv
import fitz
from docx import Document as DocxDocument
from llama_index.core import Document
import json
import time
import requests
from llama_index.core import Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import VectorStoreIndex
from llama_index.retrievers.bm25 import BM25Retriever
from llama_index.core.retrievers import VectorIndexRetriever
from HybridRetriever import HybridRetriever
from ChatEngine import ChatEngine
from configs import *

def parse_json_stream(line):
    decoded_line = line.decode('utf-8')
    decoder = json.JSONDecoder()
    pos = 0
    while pos < len(decoded_line):
        try:
            result, json_end = decoder.raw_decode(decoded_line[pos:])
            if "text" in result:
                #print(result["text"]) # debugging
                if result["text"]:
                   yield result["text"][0]
                   time.sleep(0.1)
            pos += json_end
        except json.JSONDecodeError:
            # if can't decode JSON, go next character
            pos += 1

def process_file(file):
    """
    Processes an input file based on its extension and extracts text content.
    
    The function handles:
    - **TXT**: Reads the entire content of a plain text file.
    - **CSV**: Reads and converts the content of a CSV file into a string where each row is
      joined by commas and each line is seperated by a newline character.
    - **PDF**: Extracts text from all pages of a PDF file using `fitz` library.
    - **DOCX**: Extracts text from a Word document using `python-docx` library.

    Args:
      file: The file to be processed.
    Returns:
      list[llama_index.core.Document]: A list of Document objects with the extracted text.
    Raises:
      Exception: if the file extension is not one of the supported types (txt,csv,pdf,docx).
    """

    file_extension = file.split(".")[-1].lower()

    if file_extension == 'txt':
        with open(file, 'r', encoding='utf-8') as f:
            text = f.read()

    elif file_extension == 'csv':
        with open(file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            text = '\n'.join(','.join(row) for row in reader)

    elif file_extension == 'pdf':
        pdf_document = fitz.open(file, filetype=file_extension)
        text = ""
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text")
        pdf_document.close()
        
    elif file_extension == 'docx':
        docx_document = DocxDocument(file)
        text = ""
        for paragraph in docx_document.paragraphs:
            text += paragraph.text + "\n"

    return [Document(text=text)]

def send_request_2_llm(prompt: str):
    """
    Sends a request to vllm.entrypoints.api_server with a prompt.
    
    This functions performs:
    - **Truncates** the prompt if it exceeds 27,000 characters to accommodate the model's maximum input length.
    - **Prepares** a JSON payload containing the prompt, streaming configuration, and token limits for the model.
    - **Sends** the request to the language model server running at "http://localhost:8000/generate" using a POST request.
    - **Streams** the response from the server to allow real-time processing of generated text.

    Args:
        prompt (str): The input prompt to be sent to the language model for text generation.

    Returns:
        tuple: A tuple containing:
            - `int`: The length of the prompt after potential truncation.
            - `Response`: The response object returned from the POST request to the server.
    """
    url = "http://localhost:8000/generate"
    if len(prompt) > 27_000: # model-len is set to 27k on vllm.entrypoints.api_server
        prompt = prompt[:27_000]
    payload = {
        "prompt": prompt,
        "stream": True,
        "min_tokens": 256,
        "max_tokens": 1024
    }
    response = requests.post(url, json=payload, stream=True)
    return response

def get_hybrid_retriever(file):
    """
    Processes an input file, generates a response to a given question, and returns the response length along with the server response.

    The function performs the following steps:
    - **Processes** the input file to extract text content and creates a `Document` object.
    - **Splits** the text into manageable chunks using a `SentenceSplitter` with specified chunk size and overlap.
    - **Creates** a `VectorStoreIndex` from the processed documents, embedding them with the specified model for vector-based retrieval.
    - **Initializes** two retrievers:
        - **BM25Retriever**: Uses the BM25 algorithm for similarity-based retrieval of documents.
        - **VectorIndexRetriever**: Uses the vector index for similarity-based retrieval.
    - **Combines** the two retrievers into a `HybridRetriever`, which can utilize both retrieval methods.
    - **Uses** a `ChatEngine` to generate a response to the provided question by retrieving relevant documents and combining them with the question.
    - **Sends** the generated question and relevant document text to the language model server via `send_request_2_llm`.

    Args:
        file (str): The file path to be processed.
        question (str): The question to be answered based on the file content.

    Returns:
        tuple: A tuple containing:
            - `int`: The length of the chat history (combined relevant documents and question).
            - `Response`: The response object from the language model server after processing the chat history.
    """
    documents = process_file(file)
    
    text_splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    vector_index = VectorStoreIndex.from_documents(
        documents, transformations=[text_splitter], embed_model=Settings.embed_model, show_progress=True
    )
    bm25_retriever = BM25Retriever(nodes=documents, similarity_top_k=TOP_K, tokenizer=text_splitter.split_text)
    vector_retriever = VectorIndexRetriever(index=vector_index, similarity_top_k=TOP_K)
    hybrid_retriever = HybridRetriever(bm25_retriever=bm25_retriever, vector_retriever=vector_retriever)
    return hybrid_retriever

def respond(hybrid_retriever, question):
    chat_engine = ChatEngine(hybrid_retriever)
    chat_history = chat_engine.ask_question(question)  # returns modified text with relevant documents + question
    return send_request_2_llm(chat_history)

